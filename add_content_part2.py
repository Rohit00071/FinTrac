from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document('1CR24MC085_Humanized.docx')

# 2.2 Feasibility Study
doc.add_heading('2.2 Feasibility Study', level=2)

feasibility_content = [
    'Before starting any project, it\'s important to check whether it\'s actually practical to build. I evaluated FinTrac from '
    'different angles - technical, operational, economic, and schedule feasibility - to make sure it was realistic and achievable.',
    
    '**Technical Feasibility**',
    
    'From a technical standpoint, FinTrac is definitely feasible. I used React Native for the frontend, which is reliable and '
    'widely adopted. It lets the app run smoothly on both Android and iOS. For data storage, I used local storage mechanisms '
    'that securely save user information on their device.',
    
    'All the tools and frameworks I used are open-source, so there\'s no licensing cost. They also have strong developer communities, '
    'making it easy to find documentation and solutions when I ran into issues. Since the app works efficiently on standard '
    'smartphones and web browsers, it doesn\'t require high-end hardware or specialized devices. This makes the technical '
    'implementation practical and achievable.',
    
    '**Operational Feasibility**',
    
    'Operationally, FinTrac is designed to be simple and user-friendly. The workflow matches how people naturally manage their '
    'finances. Users can easily register, log in, and access a personalized dashboard showing their financial overview. Adding '
    'transactions, setting budgets, and tracking savings goals require minimal steps and don\'t involve complex procedures.',
    
    'The app is suitable even for people with limited technical knowledge. Most actions, like recording an expense or checking '
    'budget usage, can be done with a few taps. The offline-first design also improves usability since users can manage their '
    'finances anytime without worrying about internet availability. So operationally, the system is practical and easy to adopt.',
    
    '**Economic Feasibility**',
    
    'From an economic perspective, FinTrac is cost-effective to develop and maintain. Since all development tools, frameworks, '
    'and libraries are free and open-source, there\'s no significant development cost. Local storage reduces dependency on paid '
    'cloud services, further minimizing expenses.',
    
    'The project doesn\'t require any additional hardware or external devices, making it affordable for student-level development. '
    'If I expand the app in the future, I could add optional paid services like cloud backup or advanced hosting. But for now, '
    'the overall cost remains minimal, making the project economically feasible.',
    
    '**Schedule Feasibility**',
    
    'The project is also feasible in terms of time management. I developed FinTrac in a modular way, where features like transaction '
    'management, budgeting, account handling, and goal tracking can be implemented independently. This approach ensures that delays '
    'in one module don\'t affect the entire system.',
    
    'The technologies I used are commonly taught and widely practiced, allowing faster learning and implementation. With proper '
    'planning and task division, I was able to complete the project within the allocated academic timeline. So the system is '
    'schedule-feasible and manageable within a reasonable development period.',
    
    'Overall, FinTrac is feasible from all aspects. Technically, it uses stable and accessible technologies. Operationally, it\'s '
    'easy to use and practical. Economically, it\'s low-cost. And in terms of scheduling, it can be developed within a realistic '
    'timeframe. This makes the project suitable for academic implementation and future real-world enhancement.'
]

for para in feasibility_content:
    p = doc.add_paragraph(para)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_page_break()

# 2.3 Tools and Technologies
doc.add_heading('2.3 Tools and Technologies Used', level=2)

tools_content = [
    'For developing FinTrac, I used different tools and technologies on both the frontend and backend. I selected these because '
    'they\'re free, reliable, and well-supported by the developer community. This helped ensure smooth development and made the '
    'app stable and easy to maintain.',
    
    '**Frontend Technologies**',
    
    'The frontend is what users see and interact with directly. For FinTrac, I built it using:',
    
    '**React Native** - This framework lets me build cross-platform mobile apps using JavaScript. The same codebase runs on both '
    'Android and iOS devices. I used React Native to develop screens like the dashboard, transactions page, budget management, '
    'savings goals, and account overview. One main reason I chose React Native is its component-based structure, which helps create '
    'reusable UI elements. This improves code organization and reduces development time. React Native also provides a smooth and '
    'responsive user experience without reloading screens.',
    
    '**Expo** - Expo is a tool built on top of React Native that simplifies development and testing. I used it to speed up the '
    'development process and reduce setup complexity. Expo also provides built-in features like icons and development utilities, '
    'which helped me create a clean and consistent interface.',
    
    '**Backend Technologies**',
    
    'The backend handles data processing and application logic behind the scenes. Since FinTrac follows an offline-first approach, '
    'heavy server-side processing isn\'t required. However, backend logic is handled locally within the app.',
    
    '**JavaScript (Business Logic Layer)** - I used JavaScript to implement the core business logic of the application. It manages '
    'calculations like budget usage, balance updates, goal progress tracking, and transaction handling. Using JavaScript allows '
    'better integration between frontend components and internal logic.',
    
    '**Database**',
    
    '**AsyncStorage** - I used AsyncStorage as the local database for FinTrac. It stores all user-related financial data such as '
    'accounts, transactions, budgets, and goals directly on the user\'s device. This eliminates the need for cloud storage and '
    'ensures data availability even without internet connection. Using AsyncStorage improves privacy and performance, as sensitive '
    'financial data remains under the user\'s control. It\'s lightweight, easy to integrate, and well-suited for mobile applications.',
    
    '**Authentication and Security**',
    
    'To protect user data, I implemented basic authentication features in FinTrac. Users must register and log in to access their '
    'personal dashboard and financial information. Authentication ensures that only authorized users can view or modify stored data. '
    'Sensitive data is handled securely, and access to app features is restricted based on user login status. Since data is stored '
    'locally, the risk of data leakage through external servers is minimized, enhancing overall security.',
    
    '**Development Tools**',
    
    'Apart from programming frameworks and libraries, I used several supporting tools during development:',
    
    '• **Visual Studio Code (VS Code)** – My primary code editor. It provides extensions that support React Native development.\n'
    '• **Git and GitHub** – Used for version control and code backup. These tools help track changes and manage different versions.\n'
    '• **npm** – Used to manage dependencies and install required libraries.\n'
    '• **Jest** – Used for basic testing to verify application functionality.',
    
    '**Other Libraries and Dependencies**',
    
    'Some additional libraries supported specific features in the project:',
    
    '• **UUID** – Used to generate unique IDs for transactions, accounts, budgets, and goals.\n'
    '• **Date-fns** – Used for handling dates such as transaction dates, monthly budgets, and goal deadlines.\n'
    '• **React Navigation** – Used to navigate between different screens in the application smoothly.',
    
    'The tools and technologies I used in FinTrac were chosen mainly for their simplicity, availability, and cost-effectiveness. '
    'React Native and Expo helped build a responsive and cross-platform frontend. AsyncStorage handled data storage securely and '
    'efficiently. Supporting tools like GitHub, VS Code, and npm made development more organized. Together, these tools enabled '
    'the successful development of FinTrac within the scope of a student-level project.'
]

for para in tools_content:
    p = doc.add_paragraph(para)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_page_break()

# 2.4 Hardware and Software Requirements
doc.add_heading('2.4 Hardware and Software Requirements', level=2)

requirements_content = [
    'Before developing or running any project, it\'s important to identify the required hardware and software resources. Since '
    'FinTrac is a cross-platform personal finance app designed to run on mobile devices and web browsers, the system requirements '
    'are moderate and don\'t demand high-end equipment. The app runs smoothly on standard devices with basic configurations.',
    
    '**Hardware Requirements**',
    
    '• Processor: Minimum Intel i3 or equivalent\n'
    '• RAM: 4 GB (8 GB recommended for smoother development and testing)\n'
    '• Storage: At least 500 MB of free disk space for project files and local data storage\n'
    '• Display: Standard screen resolution (1366 × 768 or higher)\n'
    '• Internet: Internet connection required during development, testing, and initial setup',
    
    'The project doesn\'t require any additional hardware like external devices, sensors, or peripherals. This makes the system '
    'easy to develop, test, and use on commonly available devices.',
    
    '**Software Requirements**',
    
    '• Operating System: Windows 10 / Linux / macOS\n'
    '• Frontend Framework: React Native with Expo\n'
    '• Programming Language: JavaScript / TypeScript\n'
    '• Local Storage: AsyncStorage\n'
    '• Code Editor: Visual Studio Code\n'
    '• Version Control: Git and GitHub\n'
    '• Testing Tool: Jest (for basic testing)\n'
    '• Browser (for web testing): Google Chrome / Firefox / Edge (latest versions recommended)',
    
    'FinTrac doesn\'t require expensive hardware or licensed software. A basic laptop or desktop system with an internet connection '
    'is sufficient for development and usage. Since all tools and frameworks used are open-source, there\'s no additional cost '
    'involved. This makes the project suitable for student-level development and allows easy scalability and enhancement in the future.'
]

for para in requirements_content:
    p = doc.add_paragraph(para)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_page_break()

doc.save('1CR24MC085_Humanized.docx')
print("Feasibility, Tools, and Requirements sections added!")
