from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document('1CR24MC085_Humanized.docx')

# 3. Software Requirements Specification
doc.add_heading('3. SOFTWARE REQUIREMENTS SPECIFICATION', level=1)

doc.add_heading('3.1 Users', level=2)

users_content = [
    'Every software system is built for specific people who will use it. In FinTrac, I identified two main types of users: '
    'guest users and registered users. In the future, I could also add an admin role if needed. Each type of user has different '
    'access levels and purposes.',
    
    '**Guest Users**',
    
    'Guest users are people who open the app without creating an account. These are usually first-time visitors who want to '
    'understand what the platform offers before signing up. Their access is intentionally limited:',
    
    '• They can view only sample or demo screens of the application\n'
    '• They can explore a basic overview of features like transaction tracking, budgeting, and goal management\n'
    '• They cannot add income or expense transactions\n'
    '• They cannot create budgets or savings goals\n'
    '• They don\'t have access to a personal dashboard or data storage',
    
    'The main purpose of guest users is to get an idea of how FinTrac works. This allows them to decide whether the app suits '
    'their needs before registering.',
    
    '**Registered Users**',
    
    'Registered users are those who have created an account and logged into the system. They receive full access to all features:',
    
    '• They can add, edit, and delete income and expense transactions\n'
    '• They can create and manage category-wise budgets and monitor spending limits\n'
    '• They can set savings goals and track progress over time\n'
    '• They can manage multiple accounts such as bank, wallet, or cash records\n'
    '• They can view a personalized dashboard that displays their financial summary\n'
    '• All their financial data is saved securely and remains private',
    
    'The main purpose of registered users is to use the app regularly as part of their daily financial routine. The dashboard '
    'provides a sense of ownership and helps users track their financial activities in an organized manner.',
    
    '**Admin Users (Optional/Future)**',
    
    'At present, FinTrac doesn\'t include an admin panel. However, in the future, an admin role could be introduced to manage '
    'and maintain the system. Possible responsibilities may include:',
    
    '• Managing application updates and maintenance\n'
    '• Monitoring system performance and stability\n'
    '• Handling user feedback and technical issues',
    
    'The admin role would not have access to individual user financial data, ensuring privacy and security.',
    
    'This user structure is kept simple to ensure ease of use and clarity. By maintaining clear and limited user roles, FinTrac '
    'remains easy to understand, efficient to use, and suitable for everyday personal finance management.'
]

for para in users_content:
    p = doc.add_paragraph(para)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_page_break()

# 3.2 Functional Requirements
doc.add_heading('3.2 Functional Requirements', level=2)

functional_content = [
    'Functional requirements describe what the system must do. These are the core features and actions that FinTrac performs '
    'when users interact with it. Without these functions, the app wouldn\'t fulfill its purpose of helping users manage their '
    'personal finances effectively.',
    
    '**1. User Authentication**',
    
    '• New users should be able to register by providing basic details like name, email, and password\n'
    '• Existing users should be able to log in using their registered credentials\n'
    '• Login credentials must be validated before granting access\n'
    '• The system should allow users to log out securely\n'
    '• If incorrect login details are entered, the system should display an appropriate error message',
    
    'This functionality is necessary to differentiate between guest users and registered users and to protect personal financial data.',
    
    '**2. Dashboard Access**',
    
    '• After successful login, users should be redirected to their personal dashboard\n'
    '• The dashboard should display a summary of the user\'s financial information\n'
    '• It should provide navigation options for transactions, budgets, goals, accounts, and reports\n'
    '• Guest users should not have access to the dashboard',
    
    '**3. Account Management**',
    
    '• Registered users should be able to add multiple financial accounts such as bank, wallet, or cash\n'
    '• Users should be able to view account balances\n'
    '• Users should be able to edit or delete accounts if required\n'
    '• The system should update balances automatically when transactions are added',
    
    '**4. Transaction Management**',
    
    '• Users should be able to add income and expense transactions\n'
    '• Each transaction should include amount, category, date, and description\n'
    '• Users should be able to view a list of all transactions\n'
    '• Users should be able to edit or delete transactions\n'
    '• Transactions should be categorized for better tracking',
    
    '**5. Budget Management**',
    
    '• Users should be able to create monthly budgets for different categories\n'
    '• The system should track spending against the set budget\n'
    '• Visual indicators should show budget usage\n'
    '• The system should alert users when spending approaches or exceeds the budget limit',
    
    '**6. Savings Goal Management**',
    
    '• Users should be able to create savings goals with target amounts\n'
    '• Users should be able to set deadlines for goals\n'
    '• The system should track progress towards each goal\n'
    '• Users should be able to add contributions to goals\n'
    '• Goal completion should be clearly indicated',
    
    '**7. Reports and Analytics**',
    
    '• The system should generate income vs expense summaries\n'
    '• Users should be able to view category-wise spending reports\n'
    '• Monthly financial trends should be displayed\n'
    '• Reports should help users understand their financial behavior',
    
    '**8. Error Handling**',
    
    '• The system should display user-friendly error messages for invalid inputs\n'
    '• Empty or incorrect fields should not be accepted\n'
    '• The application should handle errors gracefully without crashing',
    
    '**9. Logout and Session Management**',
    
    '• Users should be able to log out at any time\n'
    '• After logout, users should not be able to access the dashboard without logging in again\n'
    '• User session data should be cleared after logout to ensure security',
    
    'The functional requirements of FinTrac focus on providing a complete, simple, and secure personal finance management '
    'experience. By supporting transaction tracking, budgeting, goal setting, and reporting, the system helps users manage '
    'their finances efficiently while maintaining data privacy and ease of use.'
]

for para in functional_content:
    p = doc.add_paragraph(para)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_page_break()

# 3.3 Non-Functional Requirements
doc.add_heading('3.3 Non-Functional Requirements', level=2)

non_functional_content = [
    'While functional requirements describe what the system does, non-functional requirements explain how well the system should '
    'perform. These requirements focus on quality, performance, reliability, security, and user experience rather than specific features.',
    
    '**1. Performance Requirements**',
    
    '• The application should load within a few seconds on a normal device\n'
    '• Adding or viewing transactions should be fast and responsive\n'
    '• Navigation between dashboard screens should not lag\n'
    '• Local data operations such as saving budgets or goals should execute smoothly\n'
    '• Even with a large number of stored transactions, the system should remain stable',
    
    'Since FinTrac is meant for regular financial tracking, smooth and quick performance is important so users can record '
    'transactions without delay.',
    
    '**2. Security Requirements**',
    
    '• User login credentials must be handled securely\n'
    '• Only registered users should be able to access personal financial data\n'
    '• Guest users must not be able to view or modify stored financial information\n'
    '• After logging out, users should not be able to access the dashboard without logging in again\n'
    '• User data should remain private and protected from unauthorized access',
    
    'As FinTrac deals with sensitive financial information, security is a critical requirement of the system.',
    
    '**3. Usability Requirements**',
    
    '• The application should be easy to use without any training\n'
    '• The interface should be clean and well-organized\n'
    '• Navigation options should be clearly visible on the dashboard\n'
    '• Buttons and forms should be simple and understandable\n'
    '• The system should be usable by students, professionals, and non-technical users',
    
    'Good usability ensures that users can manage their finances comfortably and consistently.',
    
    '**4. Reliability Requirements**',
    
    '• The system should function correctly without frequent crashes\n'
    '• User data such as transactions, budgets, and goals should be saved accurately\n'
    '• In case of errors, the system should display a proper message instead of failing completely\n'
    '• Stored financial data should not be lost during normal usage',
    
    'Reliability helps users trust the application and rely on it for long-term financial tracking.',
    
    '**5. Maintainability Requirements**',
    
    '• The codebase should be well-structured and easy to understand\n'
    '• Features should be modular so updates can be added easily\n'
    '• Bug fixes should not affect existing functionality\n'
    '• Future enhancements should be simple to implement',
    
    'Since FinTrac is a student project with potential for future development, maintainability is essential.',
    
    '**6. Portability Requirements**',
    
    '• The application should run on Android, iOS, and web platforms\n'
    '• It should support major operating systems such as Windows, macOS, and Linux\n'
    '• The user interface should be responsive and adapt to different screen sizes',
    
    'Portability ensures that the system is accessible to users across multiple devices.',
    
    '**7. Scalability Requirements**',
    
    '• The system should support an increasing number of transactions per user\n'
    '• Storage should handle growing financial records over time\n'
    '• The application should allow future upgrades if usage increases\n'
    '• Additional features can be added without redesigning the entire system',
    
    'Scalability ensures that FinTrac can grow beyond a small student project and support more users if needed.',
    
    'The non-functional requirements of FinTrac ensure that the system is not only functional but also reliable, secure, '
    'user-friendly, and future-ready. These qualities are important because even a feature-rich system will fail if it\'s slow, '
    'unsafe, or difficult to use. Together, these requirements help make FinTrac a trustworthy and practical personal finance '
    'management application.'
]

for para in non_functional_content:
    p = doc.add_paragraph(para)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_page_break()

doc.save('1CR24MC085_Humanized.docx')
print("SRS sections added successfully!")
