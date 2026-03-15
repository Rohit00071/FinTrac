from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_page_break(doc):
    doc.add_page_break()

def set_heading_format(paragraph, level=1):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.runs[0]
    if level == 1:
        run.font.size = Pt(16)
        run.font.bold = True
    elif level == 2:
        run.font.size = Pt(14)
        run.font.bold = True
    else:
        run.font.size = Pt(12)
        run.font.bold = True

# Create document
doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# Title Page
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('FinTrac\n\n')
run.font.size = Pt(24)
run.font.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('A Personal Finance Management Application\n\n')
run.font.size = Pt(14)

doc.add_paragraph('\n' * 3)

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run(
    'A Dissertation submitted in partial fulfillment of the requirement\n'
    'For the award of degree of\n'
    'MASTER OF COMPUTER APPLICATIONS\n'
    'Of\n'
    'Visvesvaraya Technological University\n'
    'Belgaum, Karnataka\n\n'
    'By\n'
    'ROHIT T NADAGERI\n'
    '1CR24MC085\n\n'
    'Under the Guidance of\n'
    'Dr. KALAISELVI K\n'
    'Assistant Professor\n'
    'CMR Institute of Technology\n\n'
    'Department of Master of Computer Applications\n'
    'CMR INSTITUTE OF TECHNOLOGY\n'
    '132, IT Park Road, Kundalahalli, Bangalore – 560037\n'
    '2024-2026'
)
run.font.size = Pt(12)

add_page_break(doc)

# Certificate Page
cert_heading = doc.add_heading('CERTIFICATE', level=1)
cert_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

cert_text = doc.add_paragraph(
    'This is to certify that the project work entitled "FinTrac" submitted in partial fulfillment '
    'of the requirement for the award of the degree of Master of Computer Applications of the '
    'Visveswaraya Technological University, Belgaum, Karnataka is a bonafide work carried out by '
    'Rohit T Nadageri (1CR24MC085).'
)
cert_text.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_paragraph('\n' * 5)

signatures = doc.add_paragraph()
signatures.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = signatures.add_run(
    'Signature of Guide                    Signature of HoD                    Signature of Principal\n\n'
    'Dr. Kalaiselvi K                      Dr. Gomathi T                       Dr. Sanjay Jain\n'
    'Asst. Professor, MCA                  HoD, MCA                            Principal, CMRIT'
)

add_page_break(doc)

# Declaration
decl_heading = doc.add_heading('DECLARATION', level=1)
decl_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

decl_text = doc.add_paragraph(
    'I, ROHIT T NADAGERI, student of the 4th semester MCA at CMR Institute of Technology, '
    'Bangalore, bearing USN 1CR24MC085, hereby declare that the project entitled "FinTrac" '
    'has been carried out by me under the supervision of Prof. KALAISELVI K, Department of '
    'Computer Applications. This work is submitted in partial fulfillment of the requirements '
    'for the award of the Degree of Master of Computer Applications by the Visvesvaraya '
    'Technological University during the academic year 2024-25.\n\n'
    'This report has not been submitted to any other organization or university for any award '
    'of degree or certificate.'
)
decl_text.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_paragraph('\n' * 3)

sign = doc.add_paragraph()
sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = sign.add_run('Place: Bangalore\nDate:\n\nRohit T Nadageri\n1CR24MC085')

add_page_break(doc)

# Acknowledgement
ack_heading = doc.add_heading('ACKNOWLEDGEMENT', level=1)
ack_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

ack_text = doc.add_paragraph(
    'I would like to thank everyone who helped me complete this project.\n\n'
    'First, I want to express my sincere gratitude to my project guide, Prof. Kalaiselvi K, '
    'Assistant Professor in the MCA Department at CMR Institute of Technology. Her constant '
    'support and valuable guidance throughout this project made it possible.\n\n'
    'I also thank Dr. Gomathi, Head of Department, and Dr. Sanjay Jain, Principal of CMRIT, '
    'Bangalore, for their support during my studies.\n\n'
    'My heartfelt thanks go to my parents and family for their unconditional support.\n\n'
    'Finally, I would like to thank my friends for their advice and motivation throughout this journey.'
)
ack_text.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_paragraph('\n' * 2)

sign2 = doc.add_paragraph()
sign2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = sign2.add_run('Rohit T Nadageri\n1CR24MC085')

add_page_break(doc)

print("Creating humanized document...")
print("Part 1: Front matter completed")

# Save document
doc.save('1CR24MC085_Humanized.docx')
print("Document created successfully!")
