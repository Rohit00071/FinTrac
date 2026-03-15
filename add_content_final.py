from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document('1CR24MC085_Humanized.docx')

# 4. System Design
doc.add_heading('4. SYSTEM DESIGN', level=1)

system_design_content = [
    'System design forms the backbone of the project. It explains how different components of the system work together to deliver '
    'the required functionality. While the requirement specification phase focused on what the system should do, the design phase '
    'focuses on how the system achieves those requirements in a structured and efficient manner.',
    
    'I designed FinTrac to be modular, scalable, and easy to maintain. Each major module - transaction management, budget tracking, '
    'savings goals, and account management - operates independently but interacts through a common application structure. This '
    'modular approach allows individual features to be updated or extended without affecting the entire system.',
    
    'At a high level, FinTrac follows a three-tier architecture, which separates the system into distinct layers. This separation '
    'improves clarity, maintainability, and overall system performance.',
    
    '**Presentation Layer (Frontend)**',
    
    'The presentation layer is responsible for providing the user interface through which users interact with the app. In FinTrac, '
    'I developed this layer using React Native and Expo. It includes screens such as login and registration, dashboard, transactions, '
    'budgets, goals, and account management. This layer handles user input, displays financial information, and provides navigation '
    'between different sections of the application.',
    
    '**Application Layer (Business Logic)**',
    
    'The application layer contains the core logic of the system. It processes user actions such as adding transactions, updating '
    'budgets, calculating balances, and tracking goal progress. This layer ensures that data is validated before being stored and '
    'that business rules are followed consistently. It acts as a bridge between the user interface and the data storage layer.',
    
    '**Data Layer (Local Storage)**',
    
    'The data layer is responsible for storing and retrieving user data. In FinTrac, I handled this using local storage mechanisms '
    'such as AsyncStorage. Financial data including accounts, transactions, budgets, and goals are stored securely on the user\'s '
    'device. This approach supports offline usage and enhances data privacy, as information is not dependent on cloud servers.'
]

for para in system_design_content:
    p = doc.add_paragraph(para)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_page_break()

# 7. Software Testing
doc.add_heading('7. SOFTWARE TESTING', level=1)

testing_content = [
    'Testing is one of the most essential phases of software development. No matter how carefully an application is designed and '
    'coded, errors or bugs may still occur. Testing helps identify such issues before the software is released to real users, '
    'ensuring reliability and correctness.',
    
    'For the FinTrac project, testing played a crucial role in verifying that financial features such as Login, Transaction Management, '
    'Budget Planning, Savings Goals, and Financial Reports function correctly in real-world usage. Since FinTrac handles sensitive '
    'financial data, testing also focused on data accuracy, security, and system stability.',
    
    '**Testing Objectives**',
    
    'The main objectives of testing in FinTrac were:',
    
    '• To verify that the system meets the requirements specified in the Software Requirement Specification (SRS)\n'
    '• To ensure functional correctness - that each module works as expected\n'
    '• To test non-functional aspects such as performance, usability, and reliability\n'
    '• To confirm that the system handles both valid and invalid inputs without crashing',
    
    '**Testing Strategy**',
    
    'For the FinTrac project, I applied a combination of Black Box and White Box testing techniques:',
    
    '**Black Box Testing** - This testing focuses on system inputs and outputs without considering internal code structure. '
    'Example: Entering incorrect login credentials and verifying whether the system displays an appropriate error message.',
    
    '**White Box Testing** - This testing examines the internal logic of the application, including conditions, loops, and data '
    'handling. Example: Verifying whether transaction validation logic and data storage functions work correctly.',
    
    'Additionally, I performed the following testing levels:',
    
    '**Unit Testing** - Testing individual modules such as Login, Add Transaction, Budget Creation, and Savings Goal modules.',
    
    '**Integration Testing** - Testing combined workflows such as "Login → Add Transaction → Update Budget → View Reports".',
    
    '**System Testing** - Testing the entire application as a complete system to ensure all components work together properly.',
    
    '**Test Results**',
    
    'All major features of FinTrac were tested thoroughly. The login and authentication system correctly validated user credentials '
    'and prevented unauthorized access. Transaction management properly stored and retrieved financial records. Budget tracking '
    'accurately calculated spending against limits and displayed appropriate alerts. Savings goal management correctly tracked '
    'progress and updated completion status. The system handled invalid inputs gracefully by displaying error messages instead of '
    'crashing.',
    
    'Overall, the testing phase confirmed that FinTrac meets its functional and non-functional requirements. The app performs reliably, '
    'securely stores user data, and provides a smooth user experience across different scenarios.'
]

for para in testing_content:
    p = doc.add_paragraph(para)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_page_break()

# 8. Conclusion
doc.add_heading('8. CONCLUSION', level=1)

conclusion_content = [
    'FinTrac was developed as a practical solution to help people manage their personal finances more effectively. Throughout this '
    'project, I applied the knowledge I gained during my MCA studies to build a real-world application that addresses common financial '
    'management challenges.',
    
    'The app successfully provides essential features like income and expense tracking, budget management, savings goal setting, and '
    'multi-account handling. By focusing on simplicity and usability, FinTrac makes financial management accessible to students, '
    'professionals, and anyone who wants to take control of their money without dealing with overly complex tools.',
    
    'One of the key achievements of this project is the offline-first approach. By storing data locally on the user\'s device, FinTrac '
    'ensures that users can access and manage their financial information anytime, anywhere, without depending on internet connectivity. '
    'This also enhances privacy and gives users complete control over their sensitive financial data.',
    
    'From a technical perspective, the project demonstrates the effective use of modern technologies like React Native, Expo, and '
    'AsyncStorage to build a cross-platform mobile application. The modular architecture makes the codebase maintainable and allows '
    'for future enhancements without major restructuring.',
    
    'Through the development of FinTrac, I gained valuable hands-on experience in full-stack mobile app development, user interface '
    'design, data management, and software testing. The project taught me how to translate user requirements into functional software, '
    'handle real-world development challenges, and create solutions that people can actually use in their daily lives.',
    
    'While FinTrac currently serves as a solid foundation for personal finance management, there\'s always room for improvement and '
    'expansion. The app has the potential to grow into a more comprehensive financial platform with additional features and capabilities.',
    
    'In conclusion, FinTrac successfully achieves its goal of providing a simple, secure, and accessible personal finance management '
    'solution. It demonstrates that with the right approach and tools, we can build practical applications that make a real difference '
    'in how people manage their money and plan for their financial future.'
]

for para in conclusion_content:
    p = doc.add_paragraph(para)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_page_break()

# 9. Future Enhancements
doc.add_heading('9. FUTURE ENHANCEMENTS', level=1)

future_content = [
    'While FinTrac currently provides a solid set of features for personal finance management, there are several ways the app could '
    'be enhanced in the future to provide even more value to users:',
    
    '**1. Cloud Sync and Backup**',
    
    'Currently, all data is stored locally on the user\'s device. In the future, I could add optional cloud backup functionality '
    'that allows users to sync their data across multiple devices. This would help users access their financial information from '
    'their phone, tablet, or computer while still maintaining the offline-first approach.',
    
    '**2. AI-Powered Financial Insights**',
    
    'Integrating AI and machine learning could provide users with intelligent insights about their spending patterns. The system '
    'could analyze transaction history and offer personalized recommendations like "You\'re spending 30% more on dining out this '
    'month" or "Based on your income, you could save ₹5,000 more per month."',
    
    '**3. Bill Reminders and Recurring Transactions**',
    
    'Adding a feature to set up recurring transactions and bill reminders would help users never miss important payments. The app '
    'could automatically record monthly expenses like rent, subscriptions, and utility bills.',
    
    '**4. Investment Tracking**',
    
    'Expanding the app to include investment tracking would allow users to monitor their stocks, mutual funds, and other investments '
    'alongside their regular income and expenses. This would provide a more complete picture of their overall financial health.',
    
    '**5. Multi-Currency Support**',
    
    'While FinTrac currently focuses on Indian Rupees, adding support for multiple currencies would make the app useful for people '
    'who travel frequently or manage finances in different countries.',
    
    '**6. Data Export and Reports**',
    
    'Adding the ability to export financial data to formats like PDF, Excel, or CSV would help users create detailed reports for '
    'tax purposes or personal record-keeping.',
    
    '**7. Family/Shared Accounts**',
    
    'Implementing shared account features would allow families or couples to manage joint finances together. Users could track '
    'shared expenses, split bills, and work towards common financial goals.',
    
    '**8. Enhanced Security**',
    
    'Adding biometric authentication (fingerprint or face recognition) and data encryption would further enhance the security of '
    'sensitive financial information.',
    
    '**9. Financial Education Content**',
    
    'Including educational content about budgeting, saving, investing, and financial planning could help users make better financial '
    'decisions and improve their overall financial literacy.',
    
    '**10. Integration with Banking APIs**',
    
    'In the future, integrating with banking APIs could allow automatic transaction imports, making it even easier for users to '
    'track their expenses without manual entry.',
    
    'These enhancements would transform FinTrac from a simple personal finance tracker into a comprehensive financial management '
    'platform. However, the core principle would remain the same: keeping the app simple, user-friendly, and focused on helping '
    'people manage their money better.'
]

for para in future_content:
    p = doc.add_paragraph(para)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_page_break()

# Appendix A - User Manual
doc.add_heading('APPENDIX A - USER MANUAL', level=1)

user_manual_content = [
    '**Getting Started with FinTrac**',
    
    '**1. Installation and Setup**',
    
    '• Download and install the FinTrac app on your device\n'
    '• Open the app and you\'ll see the welcome screen\n'
    '• Click on "Sign Up" to create a new account\n'
    '• Enter your name, email, and password\n'
    '• Click "Create Account" to complete registration',
    
    '**2. Logging In**',
    
    '• Open the FinTrac app\n'
    '• Enter your registered email and password\n'
    '• Click "Login" to access your dashboard',
    
    '**3. Using the Dashboard**',
    
    'After logging in, you\'ll see your personal dashboard showing:',
    
    '• Total balance across all accounts\n'
    '• Recent transactions\n'
    '• Budget status\n'
    '• Savings goals progress\n'
    '• Quick access buttons to main features',
    
    '**4. Managing Accounts**',
    
    '• Click on "Accounts" from the dashboard\n'
    '• Click "Add Account" to create a new account\n'
    '• Enter account name (e.g., "Bank Account", "Wallet", "Cash")\n'
    '• Enter initial balance\n'
    '• Click "Save"',
    
    '**5. Adding Transactions**',
    
    '• Click on "Transactions" from the dashboard\n'
    '• Click "Add Transaction"\n'
    '• Select transaction type (Income or Expense)\n'
    '• Enter amount\n'
    '• Select category\n'
    '• Add description (optional)\n'
    '• Select date\n'
    '• Click "Save"',
    
    '**6. Creating Budgets**',
    
    '• Click on "Budgets" from the dashboard\n'
    '• Click "Create Budget"\n'
    '• Select category\n'
    '• Enter budget limit\n'
    '• Click "Save"',
    
    'The system tracks expenses against your budget and shows you how much you\'ve spent.',
    
    '**7. Setting Savings Goals**',
    
    '• Click on "Goals" from the dashboard\n'
    '• Click "Create Goal"\n'
    '• Enter goal name, target amount, and target date\n'
    '• Click "Save"',
    
    'You can add contributions to your goals anytime to track your progress.',
    
    '**8. Viewing Reports**',
    
    '• Click on "Reports" from the dashboard\n'
    '• View income vs expense summaries\n'
    '• Check category-wise spending\n'
    '• Analyze monthly trends',
    
    '**9. Troubleshooting**',
    
    '• Unable to login: Check credentials and ensure you\'re using the correct email and password\n'
    '• Transactions not showing: Refresh the page or restart the app\n'
    '• Reports not loading: Check if you have sufficient transaction data\n'
    '• For support or bug reporting, contact the support email provided on the application homepage'
]

for para in user_manual_content:
    p = doc.add_paragraph(para)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_page_break()

# Appendix B - Bibliography
doc.add_heading('APPENDIX B - BIBLIOGRAPHY', level=1)

bibliography_content = [
    '**References**',
    
    '[1] React Native Documentation - https://reactnative.dev/docs/getting-started',
    
    '[2] Expo Documentation - https://docs.expo.dev/',
    
    '[3] AsyncStorage Documentation - https://react-native-async-storage.github.io/async-storage/',
    
    '[4] React Navigation Documentation - https://reactnavigation.org/docs/getting-started',
    
    '[5] Personal Finance Management Best Practices - Various online resources and financial blogs',
    
    '[6] Mobile App Development with React Native - Online tutorials and courses',
    
    '[7] Software Engineering Principles - Course materials from MCA curriculum',
    
    '[8] UML Diagrams and System Design - Academic textbooks and online resources',
    
    '[9] Software Testing Methodologies - Course materials and industry best practices',
    
    '[10] User Interface Design Principles - Online resources and design guidelines'
]

for para in bibliography_content:
    p = doc.add_paragraph(para)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.save('1CR24MC085_Humanized.docx')
print("Final sections added! Document is complete!")
print("Saved as: 1CR24MC085_Humanized.docx")
