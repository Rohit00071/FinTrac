from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document('1CR24MC085_Humanized.docx')

def add_section(heading, content, level=1):
    h = doc.add_heading(heading, level=level)
    for para_text in content:
        p = doc.add_paragraph(para_text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return h

# Abstract
doc.add_heading('ABSTRACT', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER

abstract_content = [
    'FinTrac is a personal finance app I developed to help people track and manage their money more effectively. '
    'The app lets users record income and expenses, set budgets, create savings goals, and manage multiple accounts '
    'all in one place. I designed it specifically with Indian users in mind, so it supports rupee formatting, bank details, '
    'and UPI information.',
    
    'The app gives users a clear picture of their financial situation through organized dashboards, categorized transactions, '
    'and visual budget indicators. Users can set financial goals, track their progress, and get alerts when they\'re close '
    'to exceeding their budget limits. I built it to work offline-first, meaning you can access and manage your financial '
    'data without needing constant internet connection. This makes it more reliable and keeps your data private.',
    
    'My main goal with this project was to show how modern mobile technologies can solve real personal finance problems. '
    'FinTrac demonstrates practical use of React Native and local data storage to build a responsive, user-friendly app '
    'that works across platforms. Overall, this project applies what I learned in my studies to create something practical '
    'that helps people plan their finances better and manage their money more responsibly.'
]

for para in abstract_content:
    p = doc.add_paragraph(para)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_page_break()

# Table of Contents
toc_heading = doc.add_heading('TABLE OF CONTENTS', level=1)
toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

toc_items = [
    ('1. Introduction', '8'),
    ('    1.1 Project Description', '8'),
    ('    1.2 Company Description', '11'),
    ('2. Literature Survey', '12'),
    ('    2.1 Existing and Proposed System', '12'),
    ('    2.2 Feasibility Study', '14'),
    ('    2.3 Tools and Technologies Used', '15'),
    ('    2.4 Hardware and Software Requirements', '16'),
    ('3. Software Requirements Specification', '17'),
    ('    3.1 Users', '17'),
    ('    3.2 Functional Requirements', '19'),
    ('    3.3 Non-Functional Requirements', '21'),
    ('4. System Design', '22'),
    ('    4.1 System Perspective', '24'),
    ('    4.2 Context Diagram', '26'),
    ('5. Detailed Design', '29'),
    ('    5.1 Use Case Diagram', '29'),
    ('    5.2 Sequence Diagrams', '32'),
    ('    5.3 Collaboration Diagrams', '34'),
    ('    5.4 Activity Diagram', '36'),
    ('    5.5 Database Design', '39'),
    ('6. Implementation', '48'),
    ('    6.1 Screenshots', '48'),
    ('    6.2 Code Snippets', '50'),
    ('7. Software Testing', '54'),
    ('8. Conclusion', '55'),
    ('9. Future Enhancements', '56'),
    ('Appendix A - User Manual', '57'),
    ('Appendix B - Bibliography', '58'),
]

for item, page in toc_items:
    p = doc.add_paragraph()
    p.add_run(item).font.size = Pt(12)
    p.add_run('\t' * 5 + page).font.size = Pt(12)

doc.add_page_break()

# 1. Introduction
doc.add_heading('1. INTRODUCTION', level=1)

intro_content = [
    'Managing personal finances has become both essential and challenging in today\'s world. With multiple income sources, '
    'daily expenses, digital payments, and long-term goals, it\'s easy to lose track of where your money goes. Many existing '
    'finance apps are either too complicated, require constant internet access, or don\'t really fit how people in India manage '
    'their money. That\'s why I built FinTrac.',
    
    'FinTrac is a personal finance management app that gives users a simple, organized way to track and manage their money. '
    'Instead of juggling multiple tools or keeping manual records, users can handle everything in one place - recording income '
    'and expenses, monitoring budgets, managing savings goals, and tracking multiple accounts. The interface is clean and '
    'straightforward, so you can understand your financial situation at a glance.',
    
    'What makes FinTrac different is its focus on being user-friendly rather than overwhelming. You can add daily transactions '
    'by categorizing them into sections like food, travel, utilities, and others. Set budget limits for each category, and visual '
    'indicators show you how much you\'ve spent. This helps you spend more carefully and avoid unnecessary expenses.',
    
    'The app supports multiple financial accounts - bank accounts, wallets, and cash. You can view balances, track transactions '
    'for each account, and keep everything organized. The savings goals feature lets you set targets, deadlines, and monitor your '
    'progress over time. These features work together to help you plan for both short-term needs and long-term objectives.',
    
    'I designed FinTrac to work offline-first. This means you can access and update your financial data even without internet. '
    'All your data is stored locally on your device, which enhances privacy and gives you complete control over your information. '
    'No cloud dependency means no worries about data breaches or third-party access.',
]

for para in intro_content:
    p = doc.add_paragraph(para)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_page_break()

doc.add_heading('1.1 Project Description', level=2)

project_desc = [
    'From a technical standpoint, I built the app using React Native for the frontend, which provides a smooth cross-platform '
    'experience. Local storage mechanisms securely store financial records, ensuring fast access and reliable performance. '
    'The authentication system (login and registration) ensures only authorized users can access their data. Once logged in, '
    'users see a dashboard that acts as the central hub for navigating transactions, budgets, goals, and reports.',
    
    'FinTrac isn\'t just a tracking tool - it\'s about building better money management habits. By combining simplicity, '
    'accessibility, and practical features, the app helps users become more aware of their spending patterns, improve their '
    'financial discipline, and make informed decisions. This project demonstrates how academic knowledge can solve real-world '
    'problems and create a useful solution tailored to everyday needs.'
]

for para in project_desc:
    p = doc.add_paragraph(para)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_page_break()

print("Adding main content sections...")

# 2. Literature Survey
doc.add_heading('2. LITERATURE SURVEY', level=1)

doc.add_heading('2.1 Existing and Proposed System', level=2)

lit_survey = [
    'When it comes to managing personal finances, people already use apps like Mint, ET Money, Walnut, MoneyView, and similar '
    'tools. These apps generally offer basic features like expense tracking, budget creation, bank linking, and financial summaries. '
    'Some focus on expense visualization, while others emphasize investment tracking or bill reminders. While these systems are '
    'widely used, they have several limitations.',
    
    'One major drawback is their dependency on constant internet connectivity. Many apps store user data on cloud servers, making '
    'offline access either limited or unavailable. This is inconvenient when you want to record transactions instantly or review '
    'your finances without internet. Additionally, cloud-based storage raises privacy concerns since sensitive financial data is '
    'shared with third-party services.',
    
    'Another issue is complexity. Existing finance apps often include too many advanced features, charts, analytics, and detailed '
    'reports that can confuse average users. While these might benefit experienced users, they often overwhelm students, beginners, '
    'or people who just want to track income and expenses quickly. Some apps also require mandatory bank account linking, which '
    'not everyone is comfortable with.',
    
    'Cost is another barrier. Most platforms offer limited functionality for free and require premium subscriptions to access '
    'essential features like detailed reports, multiple accounts, or budget alerts. This makes them less suitable for students '
    'and budget-conscious users. Moreover, many apps aren\'t fully optimized for Indian financial practices, such as proper rupee '
    'formatting, lakh-crore representation, or simplified UPI-based tracking.',
    
    'To address these limitations, I designed FinTrac as a simple, user-friendly personal finance app. Instead of overwhelming '
    'users with unnecessary complexity, FinTrac focuses on essential features: income and expense tracking, category-wise budgeting, '
    'savings goal management, and multi-account handling.',
]

for para in lit_survey:
    p = doc.add_paragraph(para)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_page_break()

proposed_system = [
    'FinTrac works offline-first, allowing users to access and update their financial data without internet dependency. All data '
    'is stored locally on the user\'s device, ensuring better privacy and complete control over personal information. The app '
    'presents information through a clean dashboard, making it easy to understand your financial status at a glance.',
    
    'Another advantage is accessibility. FinTrac doesn\'t force users into paid subscriptions and is designed to be student- and '
    'beginner-friendly. The interface is straightforward, with clear navigation to transactions, budgets, goals, and reports. '
    'Indian currency formatting and localized financial structures further enhance usability for domestic users.',
    
    'In summary, while existing systems offer useful financial tools, they\'re often expensive, complex, internet-dependent, and '
    'less privacy-focused. FinTrac aims to provide a simple, secure, and accessible solution that encourages consistent financial '
    'tracking and responsible money management without unnecessary barriers.'
]

for para in proposed_system:
    p = doc.add_paragraph(para)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_page_break()

doc.save('1CR24MC085_Humanized.docx')
print("Main content sections added successfully!")
